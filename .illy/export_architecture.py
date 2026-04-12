"""Export Echo Vue Architecture Topology to DOCX format."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "Echo_Vue_Architecture_Topology.docx")

doc = Document()

# -- Page setup: landscape for diagrams --
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

# -- Styles --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    return h

def add_table_from_rows(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Light grey shading
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F0F0F0",
    })
    shading.append(shd)
    return p

# =========================================================
# TITLE PAGE
# =========================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Echo Vue by Illy Robotics")
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Production Architecture Topology & Validation")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x53, 0x34, 0x83)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Generated: April 2026  |  Document Type: Architecture Reference\nConfidential — Illy Robotics Internal")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

doc.add_page_break()

# =========================================================
# TABLE OF CONTENTS (manual)
# =========================================================
add_heading_styled("Table of Contents", level=1)
toc_items = [
    "1. Architecture Validation Summary",
    "2. Production Architecture — 3-Layer Topology",
    "   2.1 Layer 1: Edge (On-Premise Per Environment)",
    "   2.2 Layer 2: Frontend (User Devices)",
    "   2.3 Layer 3: Cloud Backend — AI MBL Engine",
    "3. Multi-Tenant Isolation — Network Binding",
    "   3.1 Hospital A vs Hospital B Scenario",
    "   3.2 Isolation Guarantees",
    "4. Data Flow — Complete Lifecycle",
    "   4.1 Setup & Calibration",
    "   4.2 Live Presence Detection",
    "5. Cloud Backend Internal Architecture",
    "6. Security Architecture",
    "7. Gaps: Current Codebase vs Production Vision",
    "8. USPTO Drawing Corrections",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# =========================================================
# 1. ARCHITECTURE VALIDATION SUMMARY
# =========================================================
add_heading_styled("1. Architecture Validation Summary", level=1)

p = doc.add_paragraph()
run = p.add_run("The Gemini-generated USPTO drawing is architecturally misleading. ")
run.bold = True
p.add_run(
    'It depicts the "AI Brain" as a local/on-premise component adjacent to the ESP32 '
    "and cameras, with no cloud backend visible. This implies all heavy AI/ML processing "
    "happens locally, which contradicts the production vision and would not scale commercially."
)

doc.add_paragraph()
add_heading_styled("USPTO Drawing Issues vs Correct Production Vision", level=2)

add_table_from_rows(
    ["Issue in USPTO Drawing", "Correct Production Vision"],
    [
        ['"AI Brain" shown as local/on-premise box', "AI MBL Engine runs in the CLOUD — GPU-intensive compute"],
        ["No cloud backend depicted", "Cloud is the central nervous system of Echo Vue"],
        ["No multi-tenant isolation shown", "Each user/org gets isolated data pipelines"],
        ["Camera/ESP32 seem to do processing", "ESP32 = dumb sensor, Camera = calibration-only then OFF"],
        ["No network binding concept", "Environments are bound to network IPs — can't manage remotely"],
        ["Monolithic single-user flow", "Multi-tenant SaaS serving homes, hospitals, businesses simultaneously"],
    ],
)

doc.add_page_break()

# =========================================================
# 2. PRODUCTION ARCHITECTURE — 3-LAYER TOPOLOGY
# =========================================================
add_heading_styled("2. Production Architecture — 3-Layer Topology", level=1)

p = doc.add_paragraph()
p.add_run(
    "Echo Vue operates across three distinct layers. No AI/ML computation occurs "
    "on the edge (ESP32) or frontenddevices. ALL heavy processing is centralized "
    "in the cloud AI MBL Engine."
)

# -- 2.1 Edge --
add_heading_styled("2.1 Layer 1: Edge (On-Premise Per Environment)", level=2)

add_code_block(
    "┌──────────────────────────────────────────────────────────┐\n"
    "│  ON-PREMISE (Hospital Floor / Home / Business)           │\n"
    "│                                                          │\n"
    "│  ┌──────────┐    WiFi CSI     ┌────────────────────┐    │\n"
    "│  │ WiFi 6   │ ──────────────► │ ESP32-S3           │    │\n"
    "│  │ Router   │   Subcarrier    │ Bridge Node        │    │\n"
    "│  │          │   Data          │                    │    │\n"
    "│  └──────────┘                 │ • CSI Extraction   │    │\n"
    "│                               │ • BLE Scanning     │    │\n"
    "│  ┌──────────┐    BLE RSSI     │ • TLS 1.3 Stream   │    │\n"
    "│  │ BLE Tags │ ──────────────► │ • NO AI/ML         │    │\n"
    "│  │ (Staff,  │                 └────────┬───────────┘    │\n"
    "│  │ Patient) │                          │ TLS 1.3        │\n"
    "│  └──────────┘                          ▼                │\n"
    "│                                ═══ TO CLOUD ═══         │\n"
    "└──────────────────────────────────────────────────────────┘"
)

doc.add_paragraph()
add_heading_styled("ESP32 Responsibilities", level=3)

add_table_from_rows(
    ["ESP32 Does NOT", "ESP32 ONLY Does"],
    [
        ["Run any ML models", "Extracts raw CSI (242 subcarriers × 2×2 MIMO × 100Hz)"],
        ["Render any images", "Passively scans BLE advertisements (RSSI + MAC)"],
        ["Store persistent data", "Streams both over TLS 1.3 to cloud"],
        ["Make decisions about presence", "Reports LED status ring for local feedback"],
    ],
)

doc.add_paragraph()

# -- 2.2 Frontend --
add_heading_styled("2.2 Layer 2: Frontend (User Devices — Same Network Required)", level=2)

add_code_block(
    "┌──────────────────────────────────────────────────────────┐\n"
    "│  USER DEVICES (Must be on same network as environment)   │\n"
    "│                                                          │\n"
    "│  ┌────────────────────────────────────────────────────┐  │\n"
    "│  │  Echo Vue Web Portal (Next.js / React)              │  │\n"
    "│  │                                                      │  │\n"
    "│  │  • Environment Setup & Floor Plan Editor             │  │\n"
    "│  │  • Calibration Wizard (temp camera for onboarding)   │  │\n"
    "│  │  • Live 3D Floor Plan Viewer                         │  │\n"
    "│  │  • Vitals Dashboard                                  │  │\n"
    "│  │  • Staff/Patient Presence Overlay                    │  │\n"
    "│  │                                                      │  │\n"
    "│  │  DOES NOT:                                           │  │\n"
    "│  │  ✗ Run AI models in production                       │  │\n"
    "│  │  ✗ Process CSI data                                  │  │\n"
    "│  │  ✗ Compute skeletal poses                            │  │\n"
    "│  │  ✗ Store other tenants' data                         │  │\n"
    "│  │                                                      │  │\n"
    "│  │  ONLY:                                               │  │\n"
    "│  │  ✓ Renders pre-computed vectors from cloud           │  │\n"
    "│  │  ✓ Displays 3D floor plans with avatars              │  │\n"
    "│  │  ✓ Shows vitals and alerts                           │  │\n"
    "│  │  ✓ Sends calibration data during setup               │  │\n"
    "│  └────────────────────────────────────────────────────┘  │\n"
    "└──────────────────────────────────────────────────────────┘"
)

doc.add_page_break()

# -- 2.3 Cloud --
add_heading_styled("2.3 Layer 3: Cloud Backend — AI MBL Engine (The Heavy Lifter)", level=2)

add_code_block(
    "╔═══════════════════════════════════════════════════════════════════╗\n"
    "║  ☁  ECHO VUE CLOUD — AI MBL ENGINE                              ║\n"
    "║                                                                   ║\n"
    "║  ┌──────────────────────────────────────────────────────────┐    ║\n"
    "║  │  API GATEWAY                                              │    ║\n"
    "║  │  • TLS Termination          • Rate Limiting               │    ║\n"
    "║  │  • JWT Validation           • DDoS Protection             │    ║\n"
    "║  │  • Network Origin Check     • Subscription Tier Enforce   │    ║\n"
    "║  └──────────────────────────────────────────────────────────┘    ║\n"
    "║                              │                                    ║\n"
    "║  ┌──────────────────────────┴──────────────────────────────┐    ║\n"
    "║  │  TENANT ISOLATION LAYER                                   │    ║\n"
    "║  │  • User UUID scoping on ALL queries                       │    ║\n"
    "║  │  • Environment ownership verification                     │    ║\n"
    "║  │  • Network IP/ID binding (env ↔ network)                  │    ║\n"
    "║  │  • Subscription tier enforcement                          │    ║\n"
    "║  └──────────────────────────┬──────────────────────────────┘    ║\n"
    "║                              │                                    ║\n"
    "║  ┌──────────────────────────┴──────────────────────────────┐    ║\n"
    "║  │  AI MBL COMPUTE ENGINE (GPU + CPU Workers)                │    ║\n"
    "║  │                                                            │    ║\n"
    "║  │  GPU Workloads:                                            │    ║\n"
    "║  │  ├─ LatentCSI Encoder (CSI → 512-dim vectors)             │    ║\n"
    "║  │  ├─ CalibrationGAN (adversarial training, 500 epochs)     │    ║\n"
    "║  │  └─ Pose Inference (3D skeletal reconstruction)           │    ║\n"
    "║  │                                                            │    ║\n"
    "║  │  CPU Workloads:                                            │    ║\n"
    "║  │  ├─ CSI Filtering (bandpass, Hampel outlier removal)      │    ║\n"
    "║  │  ├─ Kalman Tracking (6-state filter, multi-person)        │    ║\n"
    "║  │  ├─ RF Signature Engine (gait + breath + mass features)   │    ║\n"
    "║  │  └─ Vitals Processor (HR, breathing, fall detection)      │    ║\n"
    "║  └──────────────────────────┬──────────────────────────────┘    ║\n"
    "║                              │                                    ║\n"
    "║  ┌──────────────────────────┴──────────────────────────────┐    ║\n"
    "║  │  DATA LAYER                                               │    ║\n"
    "║  │  ├─ PostgreSQL (users, environments, activity logs)       │    ║\n"
    "║  │  ├─ Milvus Vector DB (RF embeddings, partitioned)         │    ║\n"
    "║  │  ├─ Redis (session state, real-time pubsub, cache)        │    ║\n"
    "║  │  └─ Object Storage (floor plans, snapshots, heatmaps)    │    ║\n"
    "║  └──────────────────────────────────────────────────────────┘    ║\n"
    "╚═══════════════════════════════════════════════════════════════════╝"
)

doc.add_page_break()

# =========================================================
# 3. MULTI-TENANT ISOLATION
# =========================================================
add_heading_styled("3. Multi-Tenant Isolation — Network Binding", level=1)

add_heading_styled("3.1 Hospital A vs Hospital B Scenario", level=2)

p = doc.add_paragraph()
p.add_run(
    "Two competing hospitals using Echo Vue will NEVER see each other's data. "
    "The AI MBL Engine processes data from multiple tenants simultaneously "
    "but returns results ONLY to the requesting tenant's authenticated session "
    "and verified network origin."
)

add_code_block(
    "Hospital A (IP: 203.0.113.10)           Hospital B (IP: 198.51.100.20)\n"
    "┌─────────────────────────┐            ┌─────────────────────────┐\n"
    "│ 12 ESP32 nodes (4/floor)│            │ 8 ESP32 nodes (2/floor) │\n"
    "│ Admin: admin@hosp-a.com │            │ Admin: admin@hosp-b.com │\n"
    "│ Env bound: 203.0.113.10 │            │ Env bound: 198.51.100.20│\n"
    "└────────────┬────────────┘            └────────────┬────────────┘\n"
    "             │                                       │\n"
    "             └──────────────┬────────────────────────┘\n"
    "                            │\n"
    "                  ┌─────────▼─────────┐\n"
    "                  │  ☁ AI MBL Engine   │\n"
    "                  │                    │\n"
    "                  │  ┌──────────────┐  │\n"
    "                  │  │Tenant Router │  │\n"
    "                  │  │JWT + Network │  │\n"
    "                  │  │Verification  │  │\n"
    "                  │  └──────┬───────┘  │\n"
    "                  │         │          │\n"
    "                  │    ┌────┴────┐     │\n"
    "                  │    ▼         ▼     │\n"
    "                  │ Pipeline A  Pipeline B\n"
    "                  │ (Hosp A     (Hosp B\n"
    "                  │  ONLY)       ONLY) │\n"
    "                  │                    │\n"
    "                  │ DB: Row-Level      │\n"
    "                  │ Security per       │\n"
    "                  │ user_id            │\n"
    "                  └────────────────────┘"
)

doc.add_paragraph()

add_heading_styled("Network Binding — Blocked Scenario", level=3)

add_table_from_rows(
    ["Scenario", "Network Check", "Result"],
    [
        [
            "Hospital A admin ON-SITE\n(IP: 203.0.113.10)",
            "203.0.113.10 = 203.0.113.10 ✓",
            "ACCESS GRANTED\nReturns Hospital A data ONLY",
        ],
        [
            "Hospital A admin at HOME\n(IP: 72.14.200.5)",
            "72.14.200.5 ≠ 203.0.113.10 ✗",
            "ACCESS DENIED\nMust be on environment's network",
        ],
        [
            "Hospital B admin ON-SITE\n(IP: 198.51.100.20)",
            "198.51.100.20 = 198.51.100.20 ✓",
            "ACCESS GRANTED\nReturns Hospital B data ONLY\nHospital A data NEVER visible",
        ],
    ],
)

doc.add_paragraph()

add_heading_styled("3.2 Isolation Guarantees", level=2)

guarantees = [
    ("Authentication", "Google OAuth → JWT with user UUID"),
    ("Authorization", "Every DB query scoped to WHERE user_id = {jwt.sub}"),
    ("Network Binding", "Environment creation records network_id (public IP); presence scans verify request origin matches"),
    ("Vector DB Isolation", "Milvus collections partitioned by environment_id — RF signatures physically separate per tenant"),
    ("Subscription Enforcement", "Tier limits (Personal: 10 envs, Pro: 50 envs) enforced at API gateway"),
]
for label, desc in guarantees:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# =========================================================
# 4. DATA FLOW
# =========================================================
add_heading_styled("4. Data Flow — Complete Lifecycle", level=1)

add_heading_styled("4.1 Setup & Calibration", level=2)

add_code_block(
    "SETUP:\n"
    "  User creates account  ─►  Google OAuth  ─►  JWT issued\n"
    "  User creates Environment  ─►  binds to current network IP\n"
    "  User registers ESP32 node  ─►  binds to environment\n"
    "\n"
    "CALIBRATION (one-time per person):\n"
    "  1. Camera ON briefly (privacy-preserving)\n"
    "  2. Frontend captures MoveNet keypoints (17-point pose)\n"
    "  3. ESP32 streams CSI simultaneously\n"
    "  4. Both sent to Cloud via WebSocket\n"
    "  5. Cloud: LatentCSI encodes CSI → 512-dim vectors\n"
    "  6. Cloud: CalibrationGAN trains (CSI ↔ pose mapping)\n"
    "  7. Cloud: RF Signature extracted (gait + breathing + mass)\n"
    "  8. Cloud: Embeddings stored in Milvus (env-partitioned)\n"
    "  9. Camera OFF permanently → RF-only mode"
)

doc.add_paragraph()
add_heading_styled("4.2 Live Presence Detection (Camera-Free)", level=2)

add_code_block(
    "LIVE PRESENCE (continuous, camera-free):\n"
    "\n"
    "  ESP32 Bridge          Cloud AI MBL Engine         Echo Vue Frontend\n"
    "  ─────────────         ──────────────────          ─────────────────\n"
    "       │                        │                          │\n"
    "       │──CSI @ 100Hz──────────►│                          │\n"
    "       │   (TLS 1.3)            │                          │\n"
    "       │                        │─ Filter CSI              │\n"
    "       │                        │─ LatentCSI decode        │\n"
    "       │                        │─ Kalman track            │\n"
    "       │                        │─ Extract vitals          │\n"
    "       │                        │─ Detect falls            │\n"
    "       │                        │                          │\n"
    "       │                        │──Tracking Snapshot──────►│\n"
    "       │                        │  (3D vectors, vitals,    │\n"
    "       │                        │   alerts, positions)     │\n"
    "       │                        │                          │─ Render 3D\n"
    "       │                        │                          │  floor plan\n"
    "       │                        │                          │─ Display\n"
    "       │                        │                          │  avatars\n"
    "       │                        │                          │─ Show vitals\n"
    "       │                        │                          │\n"
    "       │          FRONTEND DOES ZERO AI PROCESSING         │"
)

doc.add_page_break()

# =========================================================
# 5. CLOUD BACKEND INTERNAL
# =========================================================
add_heading_styled("5. Cloud Backend Internal Architecture", level=1)

add_code_block(
    "┌═══════════════════════════════════════════════════════════════════════┐\n"
    "║  CLOUD INFRASTRUCTURE (Production Deployment)                        ║\n"
    "║                                                                       ║\n"
    "║  ┌───────────────────────────────────────────────────────────────┐   ║\n"
    "║  │  LOAD BALANCER / CDN (Cloudflare / AWS ALB)                   │   ║\n"
    "║  │  TLS Termination  •  WAF  •  DDoS Shield  •  GeoDNS          │   ║\n"
    "║  └────────────────────────────┬──────────────────────────────────┘   ║\n"
    "║                               │                                       ║\n"
    "║  ┌────────────────────────────┴──────────────────────────────────┐   ║\n"
    "║  │  API PODS (Horizontal Autoscale)                               │   ║\n"
    "║  │  FastAPI Pod 1  │  FastAPI Pod 2  │  FastAPI Pod N...          │   ║\n"
    "║  └────────────────────────────┬──────────────────────────────────┘   ║\n"
    "║                               │                                       ║\n"
    "║  ┌────────────────────────────┴──────────────────────────────────┐   ║\n"
    "║  │  MESSAGE QUEUE (Kafka / RabbitMQ)                              │   ║\n"
    "║  │  CSI Frame Ingestion  •  Async Job Distribution                │   ║\n"
    "║  └──────┬─────────────────────────────────────────┬──────────────┘   ║\n"
    "║         │                                          │                  ║\n"
    "║  ┌──────▼──────────────┐   ┌──────────────────────▼──────────────┐  ║\n"
    "║  │  GPU COMPUTE PODS    │   │  CPU COMPUTE PODS                   │  ║\n"
    "║  │                      │   │                                      │  ║\n"
    "║  │  LatentCSI Encoder   │   │  Kalman Tracker (stateful/env)      │  ║\n"
    "║  │  (NVIDIA T4/A100)    │   │  CSI Filter Pipeline                │  ║\n"
    "║  │                      │   │  RF Signature Engine                 │  ║\n"
    "║  │  GAN Training Worker │   │  Vitals Processor                   │  ║\n"
    "║  │  (NVIDIA A100)       │   │  (HR, Breathing, Falls)             │  ║\n"
    "║  │                      │   │                                      │  ║\n"
    "║  │  Pose Inference      │   │                                      │  ║\n"
    "║  │  (NVIDIA T4)         │   │                                      │  ║\n"
    "║  └──────────┬───────────┘   └───────────────────────┬────────────┘  ║\n"
    "║             │                                        │                ║\n"
    "║  ┌──────────┴────────────────────────────────────────┴────────────┐  ║\n"
    "║  │  MANAGED DATA SERVICES                                          │  ║\n"
    "║  │                                                                  │  ║\n"
    "║  │  PostgreSQL (RDS)     │  Milvus Cluster      │  Redis Cluster   │  ║\n"
    "║  │  Users, Environments  │  RF Embeddings       │  Session State   │  ║\n"
    "║  │  Activity Logs        │  Partitioned/env_id  │  PubSub, Cache   │  ║\n"
    "║  │  Row-Level Security   │  COSINE similarity   │  Tracking State  │  ║\n"
    "║  │                       │                      │                   │  ║\n"
    "║  │  Object Storage (S3/GCS)                                        │  ║\n"
    "║  │  Floor Plans  •  Snapshots  •  Calibration Artifacts            │  ║\n"
    "║  └──────────────────────────────────────────────────────────────────┘  ║\n"
    "║                                                                       ║\n"
    "║  ┌──────────────────────────────────────────────────────────────────┐ ║\n"
    "║  │  OBSERVABILITY: Prometheus + Grafana + Sentry                     │ ║\n"
    "║  └──────────────────────────────────────────────────────────────────┘ ║\n"
    "╚═══════════════════════════════════════════════════════════════════════╝"
)

doc.add_page_break()

# =========================================================
# 6. SECURITY ARCHITECTURE
# =========================================================
add_heading_styled("6. Security Architecture", level=1)

add_table_from_rows(
    ["Layer", "Mechanism", "Details"],
    [
        ["Transport", "TLS 1.3", "All ESP32 → Cloud and Frontend → Cloud connections encrypted"],
        ["Authentication", "Google OAuth 2.0 + JWT", "24h expiry, HS256 signed, user UUID embedded"],
        ["Authorization", "Row-Level Security", "Every DB query scoped to authenticated user_id"],
        ["Network Binding", "IP Origin Verification", "Environment locked to creation network; presence scans verify match"],
        ["Tenant Isolation", "UUID + Environment Partitioning", "Milvus partitioned by env_id; PostgreSQL filtered by user_id"],
        ["API Protection", "Rate Limiting + WAF", "DDoS protection, request validation at gateway"],
        ["Data Separation", "Logical Isolation", "No cross-tenant queries possible; rendering returns only matched data"],
        ["Camera Privacy", "Automatic Shutoff", "Camera used only during calibration; server-side hard stop after GAN handoff"],
    ],
)

doc.add_page_break()

# =========================================================
# 7. GAPS TABLE
# =========================================================
add_heading_styled("7. Gaps: Current Codebase vs Production Vision", level=1)

add_table_from_rows(
    ["Component", "Current State", "Production Need", "Priority"],
    [
        ["Network Binding", "Not implemented", "Environment ↔ network IP binding + verification on every request", "CRITICAL"],
        ["ESP32 TLS Connection", "Protocol defined, manager is stub", "Full TLS 1.3 handler with device lifecycle", "CRITICAL"],
        ["API Gateway", "No rate limiting, no WAF", "Rate limiting, DDoS protection, request validation", "CRITICAL"],
        ["Subscription Enforcement", "Tiers defined, not enforced", "Middleware check before accepting CSI streams", "HIGH"],
        ["Redis Session State", "In-memory singleton", "Redis-backed state for horizontal scaling", "HIGH"],
        ["Real-Time PubSub", "REST polling (stale data)", "WebSocket push via Redis PubSub", "HIGH"],
        ["Cloud Storage", "Config exists, never used", "Persist floor plans, snapshots, artifacts", "HIGH"],
        ["Monitoring", "No observability stack", "Prometheus + Grafana + Sentry", "HIGH"],
        ["BLE Tethering", "Stub with placeholder logic", "Full MAC → person matching", "MEDIUM"],
        ["Federated Learning", "Flower framework stubs", "Cross-env model improvement", "MEDIUM"],
        ["Edge ML Filtering", "Firmware placeholder", "Optional pre-filtering on ESP32", "LOW"],
    ],
)

doc.add_page_break()

# =========================================================
# 8. USPTO CORRECTIONS
# =========================================================
add_heading_styled("8. What the USPTO Drawing Should Show", level=1)

p = doc.add_paragraph()
p.add_run(
    "The Gemini USPTO drawing is misleading because it depicts the system as a "
    "self-contained on-premise unit. For accurate patent drawings, the architecture "
    "should clearly show:"
)

corrections = [
    ("Three distinct zones", "Edge (ESP32), Cloud (AI MBL), Frontend (User Devices) — clearly separated with network boundaries"),
    ("Cloud as the processing center", "ALL AI/ML computation happens in the cloud, not on-premise"),
    ("Network arrows", "TLS 1.3 from edge → cloud, HTTPS/WSS from frontend → cloud — showing internet traversal"),
    ("Multi-tenancy", "Multiple environments served by one cloud engine with tenant isolation boundaries"),
    ("Network binding", "Dashed line showing environment-to-network-IP binding constraint"),
    ("Data return path", 'Cloud → frontend with "rendered vectors only" annotation — no raw data leaves the cloud'),
    ("Camera lifecycle", "Clearly show camera ON during calibration → OFF during live mode (server-enforced)"),
]
for label, desc in corrections:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(desc)

# =========================================================
# PROCESSING DISTRIBUTION TABLE (Appendix)
# =========================================================
doc.add_page_break()
add_heading_styled("Appendix: Processing Distribution by Layer", level=1)

add_table_from_rows(
    ["Task", "Runs On", "Reason"],
    [
        ["CSI Extraction", "Edge (ESP32)", "Must stream raw 242 subcarriers × 100Hz"],
        ["BLE Scanning", "Edge (ESP32)", "Passive advertisement collection"],
        ["CSI Filtering (bandpass, outlier)", "Cloud Backend", "Requires full history window; stateful"],
        ["LatentCSI Encoding", "Cloud Backend (GPU)", "Torch model; requires GPU/VRAM"],
        ["GAN Calibration Training", "Cloud Backend (GPU)", "500-epoch training loop; requires GPU"],
        ["Pose Inference (CSI-only)", "Cloud Backend (GPU)", "LatentCSI decoder; one-shot per frame"],
        ["Kalman Tracking", "Cloud Backend (CPU)", "Stateful per-track; centralized multi-device state"],
        ["RF Signature Extraction", "Cloud Backend (CPU)", "Gait + breathing + mass feature engineering"],
        ["Vitals Processing", "Cloud Backend (CPU)", "Heart rate, breathing, fall detection"],
        ["MoveNet Pose Detection", "Frontend (calibration only)", "Temporary label collection; TensorFlow.js"],
        ["3D Floor Plan Rendering", "Frontend", "Displays pre-computed vectors from cloud"],
        ["User Settings Sync", "Frontend + Cloud", "Frontend debounces writes; cloud is CRUD store"],
    ],
)

# =========================================================
# SAVE
# =========================================================
doc.save(OUT_PATH)
print(f"✅ Exported to: {OUT_PATH}")
